from __future__ import annotations
"""Extractor de texto para documentos Word (.docx) usando python-docx."""

from pathlib import Path

import docx
import structlog

from jurisbot.ingestion.extractors.base import BaseExtractor, ExtractionResult
from jurisbot.exceptions import ExtractionError

logger = structlog.get_logger()


class DocxExtractor(BaseExtractor):
    """Extrae texto de documentos Word preservando estructura y estilos."""

    def supported_formats(self) -> list[str]:
        return ["DOCX"]

    def extract(self, file_path: Path) -> ExtractionResult:
        """Extrae texto de un archivo .docx preservando estructura.

        Args:
            file_path: Ruta al archivo Word.

        Returns:
            ExtractionResult con texto, estructura y tablas.
        """
        try:
            document = docx.Document(str(file_path))
        except Exception as e:
            raise ExtractionError(
                f"No se pudo abrir el archivo Word '{file_path.name}': {e}"
            ) from e

        paragraphs_text: list[str] = []
        structured_sections: list[dict] = []
        tables: list[dict] = []

        # Extraer metadatos
        core = document.core_properties
        metadata = {
            "title": core.title or "",
            "author": core.author or "",
            "subject": core.subject or "",
            "created": str(core.created) if core.created else "",
            "modified": str(core.modified) if core.modified else "",
            "format": "DOCX",
            "extractor": "python-docx",
        }

        # Extraer párrafos con detección de estructura
        for i, para in enumerate(document.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            paragraphs_text.append(text)

            # Detectar headings
            style_name = (para.style.name or "").lower()
            if "heading" in style_name or "título" in style_name:
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1

                structured_sections.append({
                    "type": "heading",
                    "level": level,
                    "text": text,
                    "paragraph_index": i,
                })

            # Detectar texto en negrita como posible sub-encabezado
            elif para.runs and all(run.bold for run in para.runs if run.text.strip()):
                if len(text) < 150:  # Probablemente un título
                    structured_sections.append({
                        "type": "bold_heading",
                        "text": text,
                        "paragraph_index": i,
                    })

        # Extraer tablas
        for table_idx, table in enumerate(document.tables):
            rows_data = []
            headers = []
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_idx == 0:
                    headers = cells
                else:
                    rows_data.append(cells)

            if headers or rows_data:
                tables.append({
                    "table_index": table_idx,
                    "headers": headers,
                    "rows": rows_data,
                })

        raw_text = "\n\n".join(paragraphs_text)
        metadata["pages"] = max(1, len(raw_text) // 3000)  # Estimación

        logger.info(
            "docx_extracted",
            file=file_path.name,
            paragraphs=len(paragraphs_text),
            sections=len(structured_sections),
            tables=len(tables),
            chars=len(raw_text),
        )

        return ExtractionResult(
            raw_text=raw_text,
            pages=metadata["pages"],
            metadata=metadata,
            structured_sections=structured_sections,
            tables=tables,
            confidence_score=0.95,
        )
